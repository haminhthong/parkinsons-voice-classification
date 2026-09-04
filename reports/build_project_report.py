# ruff: noqa: E501
"""Tạo báo cáo phân tích dự án dạng DOCX bằng python-docx."""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "reports" / "Bao_cao_phan_tich_du_an_Parkinson_Voice.docx"
FIGURES = ROOT / "reports" / "figures"
ARTIFACTS = ROOT / "artifacts"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = "5F6B76"
GREEN = "207A4A"
GOLD = "8A6400"
RED = "9B1C1C"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths: list[int]) -> None:
    """Đặt hình học bảng theo DXA, tổng chiều rộng 9360."""
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for cell, text in zip(header.cells, headers, strict=True):
        set_cell_shading(cell, LIGHT_GRAY)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(NAVY)
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values, strict=True):
            cell.text = str(value)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc, label: str, text: str, *, color=BLUE, fill="F4F6F9") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    lead = p.add_run(f"{label}: ")
    lead.bold = True
    lead.font.color.rgb = RGBColor.from_string(color)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc, filename: str, caption: str, width=6.2) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    shape = paragraph.add_run().add_picture(str(FIGURES / filename), width=Inches(width))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", filename)
    cap = doc.add_paragraph(caption, style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    run.font.size = Pt(9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("222222")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)

    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = "PARKINSON VOICE CLASSIFICATION  |  BÁO CÁO PHÂN TÍCH KỸ THUẬT"
        header.style = styles["Header"]
        header.runs[0].font.size = Pt(8.5)
        header.runs[0].font.bold = True
        header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
        add_page_number(section.footer.paragraphs[0])


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    kicker = doc.add_paragraph("BÁO CÁO PHÂN TÍCH DỰ ÁN DỮ LIỆU")
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.runs[0]
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(14)
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Leakage-Aware Parkinson’s\nVoice Classification")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    subtitle = doc.add_paragraph(
        "Đánh giá theo bệnh nhân, hiệu chỉnh xác suất và ứng dụng demo phục vụ nghiên cứu"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor.from_string(DARK_BLUE)
    subtitle.paragraph_format.space_after = Pt(34)

    add_callout(
        doc,
        "Thông điệp trung tâm",
        "Giá trị nổi bật của dự án không nằm ở Accuracy cao, mà ở việc phát hiện phép chia "
        "ngẫu nhiên làm rò rỉ bệnh nhân và xây dựng lại toàn bộ quy trình đánh giá đáng tin cậy.",
        color=GREEN,
        fill="EAF5EF",
    )
    metadata = doc.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.space_before = Pt(34)
    metadata.add_run(f"Phiên bản báo cáo: {date.today().strftime('%d/%m/%Y')}\n").bold = True
    metadata.add_run("Phạm vi: mã nguồn, mô hình, kiểm thử, triển khai demo và định hướng CV")
    doc.add_page_break()


def add_executive_summary(doc: Document, metrics: dict) -> None:
    doc.add_heading("1. Tóm tắt điều hành", level=1)
    doc.add_paragraph(
        "Dự án giải quyết bài toán phân loại bệnh Parkinson từ 22 đặc trưng giọng nói. "
        "Mỗi bệnh nhân có nhiều bản ghi, vì vậy đơn vị độc lập thực sự là bệnh nhân chứ không "
        "phải từng dòng dữ liệu. Phiên bản hiện tại đã chuyển từ notebook thử nghiệm sang "
        "repository tái lập, có pipeline huấn luyện, đánh giá theo nhóm, API, ứng dụng Streamlit, "
        "artifact mô hình, biểu đồ và bộ test tự động."
    )
    add_table(
        doc,
        ["Hạng mục", "Kết luận kiểm tra"],
        [
            [
                "Rò rỉ dữ liệu",
                "Train/test và mọi fold CV tách theo subject_id; có assertion chống overlap.",
            ],
            ["Benchmark", "6 mô hình, gồm Dummy baseline; fold validation luôn có cả hai lớp."],
            ["Mô hình triển khai", metrics["champion"]],
            ["Quy tắc bệnh nhân", f"Gộp max; ngưỡng OOF = {metrics['decision_threshold']:.3f}."],
            ["Kiểm thử", "19/19 test đạt; Ruff không còn lỗi."],
            ["Phạm vi sử dụng", "Chỉ nghiên cứu/học tập; không dùng để chẩn đoán."],
        ],
        [2600, 6760],
    )
    add_callout(
        doc,
        "Kết luận",
        "Repository đã đủ tốt để trình bày trong CV ở vai trò dự án ML có tư duy đánh giá. "
        "Không nên quảng bá 87,5% Accuracy như bằng chứng lâm sàng vì holdout chỉ có 8 bệnh nhân "
        "và khoảng tin cậy còn rất rộng.",
        color=GOLD,
        fill="FFF8E8",
    )


def add_problem_and_data(doc: Document) -> None:
    doc.add_heading("2. Bài toán và dữ liệu", level=1)
    doc.add_heading("2.1 Đơn vị dự đoán", level=2)
    doc.add_paragraph(
        "CSV gốc có 195 bản ghi giọng nói thuộc 32 bệnh nhân. Cột name nhận diện bản ghi và "
        "được chuyển thành subject_id; status là nhãn nhị phân 0/1. Mô hình dự đoán từng bản ghi, "
        "sau đó gộp xác suất thành một kết quả cho mỗi bệnh nhân."
    )
    add_table(
        doc,
        ["Thành phần", "Quy mô / vai trò", "Kiểm soát"],
        [
            ["Bản ghi", "195 dòng", "Kiểm tra kiểu số, giá trị thiếu và schema."],
            ["Bệnh nhân", "32 subject_id", "Đơn vị chia train/test, CV và bootstrap."],
            ["Đặc trưng", "22 biến âm học gốc", "Danh sách cố định ORIGINAL_FEATURES."],
            ["Nhãn", "status ∈ {0,1}", "Từ chối nhãn ngoài miền."],
        ],
        [1900, 2500, 4960],
    )
    doc.add_heading("2.2 Vì sao chia ngẫu nhiên theo dòng là sai", level=2)
    doc.add_paragraph(
        "Các bản ghi của cùng một người có cấu trúc giọng nói tương quan. Nếu một bản ghi của "
        "bệnh nhân nằm trong train và bản ghi khác của chính người đó nằm trong test, mô hình có "
        "thể nhận diện người nói thay vì học tín hiệu liên quan đến bệnh. Accuracy 97,44% của phép "
        "audit tái lập đạt 92,31% nhưng 24/24 bệnh nhân test cũng xuất hiện trong train."
    )
    add_callout(
        doc,
        "Bài học phỏng vấn",
        "Khi dữ liệu có nhiều quan sát trên cùng một cá thể, khóa thiết kế đánh giá trước khi tối ưu "
        "mô hình. Group leakage có thể làm hỏng toàn bộ kết luận dù code chạy không lỗi.",
        color=RED,
        fill="FCEEEE",
    )


def add_methodology(doc: Document) -> None:
    doc.add_heading("3. Phương pháp leakage-aware", level=1)
    doc.add_heading("3.1 Chia dữ liệu và cross-validation", level=2)
    doc.add_paragraph(
        "Quy trình tạo bảng một dòng cho mỗi subject_id, dùng StratifiedKFold trên bảng bệnh nhân, "
        "rồi ánh xạ subject train/validation trở lại các bản ghi. Mỗi fold được kiểm tra bằng hai "
        "điều kiện bắt buộc: tập subject không giao nhau và y_valid có đủ hai lớp. Cách này sửa lỗi "
        "fold validation chỉ có status=1 từng làm ROC-AUC thành NaN."
    )
    doc.add_heading("3.2 Tiền xử lý và lựa chọn mô hình", level=2)
    doc.add_paragraph(
        "StandardScaler và SelectKBest nằm bên trong scikit-learn Pipeline, do đó mỗi fold chỉ fit "
        "tiền xử lý trên phần train. Benchmark gồm Logistic Regression, KNN, SVM-RBF dùng decision "
        "score, Random Forest, HistGradientBoosting và Dummy. SVM không dùng probability=True trong "
        "benchmark, tránh calibration nội bộ không biết subject_id."
    )
    add_figure(
        doc, "model_benchmark.png", "Hình 1. F1-macro CV theo bệnh nhân của sáu mô hình benchmark."
    )
    doc.add_heading("3.3 Hiệu chỉnh và quy tắc ra quyết định", level=2)
    doc.add_paragraph(
        "Champion KNN được hiệu chỉnh sigmoid bằng các fold theo bệnh nhân. Trên dự đoán OOF của "
        "train, dự án đồng thời chọn cách gộp xác suất (mean/median/max) và ngưỡng quyết định, với "
        "ràng buộc Specificity tối thiểu 0,5. Quy tắc khóa được là max và ngưỡng 0,835; holdout không "
        "tham gia lựa chọn champion, calibration, aggregation hay threshold."
    )
    add_figure(
        doc, "threshold_aggregation.png", "Hình 2. Tìm kiếm quy tắc gộp và ngưỡng trên OOF train."
    )


def add_results(doc: Document, metrics: dict, ci: pd.DataFrame) -> None:
    doc.add_heading("4. Kết quả và diễn giải", level=1)
    oof = metrics["oof_calibration"]
    holdout = metrics["holdout_subject"]
    rows = []
    for key in [
        "Accuracy",
        "Balanced Accuracy",
        "Recall/Sensitivity",
        "Specificity",
        "F1-macro",
        "ROC-AUC",
        "Brier score",
    ]:
        rows.append([key, f"{oof[key]:.3f}", f"{holdout[key]:.3f}"])
    add_table(doc, ["Chỉ số", "OOF train", "Holdout bệnh nhân"], rows, [3600, 2880, 2880])
    doc.add_paragraph(
        "Holdout đạt Accuracy 0,875, Recall 1,0 và ROC-AUC 0,75. Tuy nhiên Specificity chỉ 0,5, "
        "nghĩa là một trong hai bệnh nhân âm tính bị dự đoán dương tính. OOF F1-macro 0,365 và "
        "Recall 0,167 thấp hơn nhiều, cho thấy mô hình không ổn định khi số bệnh nhân huấn luyện nhỏ. "
        "Brier score khoảng 0,18 và ECE khoảng 0,12 cho thấy xác suất đã được hiệu chỉnh nhưng chưa đủ "
        "cơ sở để diễn giải như nguy cơ lâm sàng."
    )
    add_figure(
        doc, "holdout_probabilities.png", "Hình 3. Xác suất status=1 của từng bệnh nhân holdout."
    )
    doc.add_heading("4.1 Khoảng tin cậy bootstrap", level=2)
    selected = ci[
        ci["Metric"].isin(["Accuracy", "Balanced Accuracy", "Specificity", "F1-macro", "ROC-AUC"])
    ]
    ci_rows = [
        [
            row["Metric"],
            f"{float(row['Point estimate']):.3f}",
            f"[{float(row['CI 2.5%']):.3f}; {float(row['CI 97.5%']):.3f}]",
        ]
        for _, row in selected.iterrows()
    ]
    add_table(doc, ["Chỉ số", "Ước lượng", "95% CI theo bệnh nhân"], ci_rows, [3600, 2200, 3560])
    add_callout(
        doc,
        "Diễn giải thận trọng",
        "CI của Accuracy là [0,625; 1,000], Balanced Accuracy là [0,500; 1,000] và Specificity là "
        "[0,000; 1,000]. Độ rộng này phản ánh bất định do holdout chỉ có 8 bệnh nhân, không phải lỗi "
        "của bootstrap.",
        color=GOLD,
        fill="FFF8E8",
    )


def add_feature_and_error_analysis(doc: Document) -> None:
    doc.add_heading("5. Đặc trưng và phân tích lỗi", level=1)
    doc.add_paragraph(
        "SelectKBest được fit lại trong từng fold. Biểu đồ độ ổn định cho biết đặc trưng nào thường "
        "được chọn, hữu ích để đánh giá tính bền vững của tín hiệu. Đây chỉ là thống kê mô tả của mô "
        "hình, không chứng minh quan hệ nhân quả sinh học."
    )
    add_figure(
        doc, "feature_selection_stability.png", "Hình 4. Tần suất đặc trưng được chọn qua các fold."
    )
    doc.add_paragraph(
        "Ở holdout có 7/8 bệnh nhân được phân loại đúng, một false positive và không có false negative. "
        "False positive cần được xem như tín hiệu để kiểm tra ngưỡng, chất lượng thu âm, sai lệch theo "
        "thiết bị và tính đại diện của nhóm âm tính; không nên giải thích một đặc trưng đơn lẻ như nguyên "
        "nhân gây bệnh."
    )


def add_architecture(doc: Document) -> None:
    doc.add_heading("6. Kiến trúc repository và luồng triển khai", level=1)
    add_table(
        doc,
        ["Mô-đun", "Trách nhiệm chính"],
        [
            [
                "src/data.py",
                "Đọc CSV, kiểm tra schema, tạo subject_id, chia holdout theo bệnh nhân.",
            ],
            ["src/features.py", "Tạo Pipeline tiền xử lý và mô hình."],
            [
                "src/evaluate.py",
                "Fold theo bệnh nhân, metric, aggregation, threshold và bootstrap.",
            ],
            ["src/train.py", "Benchmark, chọn champion, hiệu chỉnh và lưu artifact."],
            ["src/predict.py", "Nạp artifact, dự đoán bản ghi và gộp kết quả bệnh nhân."],
            ["src/report.py", "Tái tạo bốn biểu đồ portfolio bằng backend headless."],
            ["app/api.py", "REST API: health và upload CSV để dự đoán."],
            ["app/streamlit_app.py", "Demo trực quan, bảng, biểu đồ, tải CSV kết quả."],
        ],
        [2600, 6760],
    )
    doc.add_heading("6.1 Luồng ứng dụng demo", level=2)
    doc.add_paragraph(
        "Người dùng tải CSV → validate_dataframe kiểm tra name và 22 đặc trưng → artifact Pipeline "
        "dự đoán xác suất cho từng bản ghi → normalize_aggregation áp dụng quy tắc max → so sánh với "
        "ngưỡng 0,835 → hiển thị kết quả bản ghi, kết quả bệnh nhân, biểu đồ đặc trưng và nút tải CSV."
    )
    add_callout(
        doc,
        "Cảnh báo sản phẩm",
        "FastAPI và Streamlit cùng dùng một hằng số cảnh báo nghiên cứu. Ứng dụng không yêu cầu nhập "
        "tay 22 đặc trưng và từ chối CSV thiếu cột.",
        color=RED,
        fill="FCEEEE",
    )


def add_code_audit(doc: Document) -> None:
    doc.add_heading("7. Audit code và phần đã làm sạch", level=1)
    doc.add_paragraph(
        "Audit tập trung vào trùng lặp có rủi ro làm hai nhánh hành vi lệch nhau, file giữ chỗ đã hết "
        "vai trò và lỗi chỉ xuất hiện trong CI/headless. Không xóa notebook, model card hay các bảng "
        "artifact vì chúng có mục đích riêng và còn được README/test sử dụng."
    )
    add_table(
        doc,
        ["Phát hiện", "Thay đổi", "Lợi ích"],
        [
            [
                "Hai cách lấy xác suất lớp dương",
                "Gom vào positive_class_probability trong src/utils.py.",
                "Một nguồn xử lý classes_ và kiểm tra [0,1].",
            ],
            [
                "Hai bảng ánh xạ mean/median/max",
                "Gom vào normalize_aggregation.",
                "Giảm lệch giữa evaluate và predict.",
            ],
            [
                "Đường dẫn artifact và cảnh báo lặp",
                "Tạo app/settings.py dùng chung.",
                "API/Streamlit nhất quán; chạy độc lập cwd.",
            ],
            [
                "Matplotlib gọi Tk trong test",
                "Đặt backend Agg trước khi import pyplot.",
                "Chạy được trên CI, Docker và server headless.",
            ],
            [
                "Hai file .gitkeep dư thừa",
                "Xóa khỏi artifacts và reports/figures.",
                "Hai thư mục đã có file thực, không cần giữ chỗ.",
            ],
            [
                "Artifact sai phiên bản",
                "Huấn luyện lại bằng scikit-learn 1.5.2.",
                "Loại cảnh báo unpickle 1.9.0 → 1.5.2.",
            ],
        ],
        [2500, 3660, 3200],
    )
    add_callout(
        doc,
        "Khả năng khôi phục",
        "Hai file .gitkeep đã xóa là file giữ chỗ rỗng và có thể khôi phục từ Git nếu cần. Không có dữ "
        "liệu, notebook hay mã nghiệp vụ nào bị xóa.",
        color=GREEN,
        fill="EAF5EF",
    )


def add_quality_and_limits(doc: Document) -> None:
    doc.add_heading("8. Chất lượng, kiểm thử và giới hạn", level=1)
    add_table(
        doc,
        ["Nhóm test", "Bằng chứng"],
        [
            ["Dữ liệu", "Đúng 22 đặc trưng; status chỉ 0/1; thiếu cột inference bị từ chối."],
            ["Chống leakage", "Không overlap train/test; mọi validation fold có đủ hai lớp."],
            ["Pipeline", "Scaler chỉ tồn tại trong Pipeline."],
            ["Artifact", "Load lại và tái tạo dự đoán; metadata calibration/OOF tồn tại."],
            ["Xác suất", "Mọi giá trị nằm trong [0,1]."],
            ["API và báo cáo", "Upload CSV hoạt động; bốn biểu đồ tái tạo được ở chế độ headless."],
        ],
        [2600, 6760],
    )
    doc.add_paragraph(
        "Kết quả xác minh cuối: Ruff — All checks passed; Pytest — 19 passed. Còn 14 cảnh báo deprecation từ phụ thuộc Matplotlib/PyParsing, không phải lỗi logic và không làm test thất bại."
    )
    doc.add_heading("8.1 Giới hạn bắt buộc phải công bố", level=2)
    doc.add_paragraph(
        "Mẫu chỉ gồm 32 bệnh nhân, phân bố lớp không cân bằng và xuất phát từ một bộ dữ liệu nghiên "
        "cứu nhỏ. Không có external validation, metadata thiết bị/điều kiện thu âm đầy đủ, fairness "
        "audit hay đánh giá drift. Một holdout duy nhất không đại diện cho hiệu năng ngoài thực tế. "
        "Ứng dụng vì thế chỉ là demo kỹ thuật, không phải thiết bị y tế hoặc công cụ sàng lọc."
    )


def add_cv_story_and_roadmap(doc: Document) -> None:
    doc.add_heading("9. Cách kể dự án trong CV và phỏng vấn", level=1)
    add_callout(
        doc,
        "Mô tả CV đề xuất",
        "Phát hiện Accuracy 92,31% bị lạc quan do 24/24 bệnh nhân test xuất hiện trong train; thiết "
        "kế lại đánh giá theo bệnh nhân bằng stratified subject-level CV, đưa tiền xử lý vào Pipeline, "
        "hiệu chỉnh xác suất theo nhóm và triển khai demo Streamlit/FastAPI với 19 test tự động.",
        color=GREEN,
        fill="EAF5EF",
    )
    doc.add_heading("9.1 Điểm nên nhấn mạnh khi phỏng vấn", level=2)
    doc.add_paragraph(
        "Trình bày theo chuỗi: phát hiện bất thường → xác định unit of independence → sửa split → khóa "
        "quy trình lựa chọn trên OOF train → đánh giá holdout và CI → công bố giới hạn. Nếu được hỏi vì "
        "sao không chọn mô hình có điểm test cao nhất, giải thích rằng test không được dùng để chọn "
        "champion và performance đáng tin cậy quan trọng hơn một con số đẹp."
    )
    doc.add_heading("9.2 Lộ trình cải thiện ưu tiên", level=2)
    add_table(
        doc,
        ["Ưu tiên", "Cải thiện", "Tiêu chí hoàn thành"],
        [
            [
                "P0",
                "Bổ sung dữ liệu bệnh nhân độc lập/external cohort.",
                "CI hẹp hơn; metric ổn định giữa cohort.",
            ],
            [
                "P0",
                "Nested CV theo bệnh nhân cho tuning và ước lượng tổng quát hóa.",
                "Không dùng cùng OOF để vừa chọn vừa báo cáo.",
            ],
            [
                "P1",
                "Đánh giá calibration curve, decision-curve và chi phí FP/FN.",
                "Ngưỡng gắn với mục tiêu nghiên cứu rõ ràng.",
            ],
            [
                "P1",
                "Schema versioning và metadata artifact đầy đủ.",
                "Kiểm tra phiên bản sklearn/data trước khi load.",
            ],
            [
                "P2",
                "CI Docker, smoke test API và kiểm tra ảnh regression.",
                "Một lệnh tái tạo app, artifact, test và figures.",
            ],
            [
                "P2",
                "Giám sát drift/fairness khi có metadata hợp lệ.",
                "Báo cáo theo nhóm và cảnh báo dịch chuyển dữ liệu.",
            ],
        ],
        [1200, 4960, 3200],
    )


def add_appendix(doc: Document, benchmark: pd.DataFrame) -> None:
    doc.add_heading("Phụ lục A. Benchmark đầy đủ", level=1)
    rows = [
        [
            row["Model"],
            f"{float(row['CV F1-macro mean']):.3f}",
            f"{float(row['CV F1-macro std']):.3f}",
            f"{float(row['CV Balanced Accuracy']):.3f}",
            f"{float(row['CV ROC-AUC']):.3f}",
        ]
        for _, row in benchmark.iterrows()
    ]
    add_table(
        doc, ["Mô hình", "F1", "SD", "Bal. Acc.", "ROC-AUC"], rows, [3600, 1300, 1300, 1580, 1580]
    )
    doc.add_heading("Phụ lục B. Kết luận nghiệm thu", level=1)
    doc.add_paragraph(
        "Sau đợt cải thiện: code dùng chú thích tiếng Việt nhất quán; logic dùng chung đã được tách; "
        "file dư thừa được loại bỏ; artifact và biểu đồ được tái tạo; lint và 19 test đều đạt. Dự án "
        "đủ hoàn chỉnh để đưa lên GitHub/CV như một case study về data leakage, grouped data, model "
        "evaluation, reproducibility và triển khai demo có giới hạn sử dụng rõ ràng."
    )


def build() -> Path:
    metrics = json.loads((ARTIFACTS / "metrics.json").read_text(encoding="utf-8"))
    benchmark = pd.read_csv(ARTIFACTS / "model_benchmark.csv")
    ci = pd.read_csv(ARTIFACTS / "holdout_bootstrap_ci.csv")

    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    add_executive_summary(doc, metrics)
    add_problem_and_data(doc)
    add_methodology(doc)
    add_results(doc, metrics, ci)
    add_feature_and_error_analysis(doc)
    add_architecture(doc)
    add_code_audit(doc)
    add_quality_and_limits(doc)
    add_cv_story_and_roadmap(doc)
    add_appendix(doc, benchmark)

    properties = doc.core_properties
    properties.title = "Báo cáo phân tích dự án Leakage-Aware Parkinson’s Voice Classification"
    properties.subject = "Phân tích kỹ thuật, audit code và định hướng CV"
    properties.author = "Parkinson Voice Classification Project"
    properties.keywords = "Parkinson, voice, machine learning, data leakage, grouped CV"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
